# import os
# import base64
# import json
# from datetime import datetime
# from dotenv import load_dotenv
# from flask import Flask, request, jsonify, Response, stream_with_context
# from flask_cors import CORS
# from groq import Groq
# import io
# import pytesseract
# from PyPDF2 import PdfReader
# from PIL import Image
# from docx import Document
# import openpyxl
# import csv
# from pathlib import Path
# import hashlib

# # Load environment variables
# load_dotenv()

# # Check for required environment variables
# GROQ_API_KEY = os.environ.get("GROQ_API_KEY")
# if not GROQ_API_KEY:
#     raise ValueError("GROQ_API_KEY environment variable is not set. Please set it in your .env file.")

# app = Flask(__name__)
# CORS(app, resources={
#     r"/*": {
#         "origins": [
#             "http://localhost:3000",  # React default port
#             "http://localhost:5173",  # Vite default port
#             "http://127.0.0.1:3000",
#             "http://127.0.0.1:5173",
#             "https://edu-hub-v1-f2ek.vercel.app"
#         ],
#         "methods": ["GET", "POST", "PUT", "DELETE", "OPTIONS"],
#         "allow_headers": ["Content-Type", "Authorization"]
#     }
# })

# # Initialize Groq client with error handling
# try:
#     client = Groq(api_key=GROQ_API_KEY)
# except Exception as e:
#     print(f"Error initializing Groq client: {str(e)}")
#     raise

# # Store conversation context and uploaded files
# conversation_contexts = {}
# file_contexts = {}

# class FileProcessor:
#     """Enhanced file processor supporting multiple formats"""
    
#     @staticmethod
#     def extract_text_from_pdf(file_stream):
#         """Enhanced PDF text extraction with better formatting"""
#         try:
#             reader = PdfReader(file_stream)
#             text = ""
#             metadata = {
#                 'total_pages': len(reader.pages),
#                 'title': reader.metadata.get('/Title', 'Unknown') if reader.metadata else 'Unknown',
#                 'author': reader.metadata.get('/Author', 'Unknown') if reader.metadata else 'Unknown'
#             }
            
#             for i, page in enumerate(reader.pages):
#                 page_text = page.extract_text() or ""
#                 if page_text.strip():
#                     text += f"\n--- Page {i+1} ---\n{page_text}\n"
            
#             return text, metadata
#         except Exception as e:
#             raise Exception(f"PDF processing error: {str(e)}")
    
#     @staticmethod
#     def extract_text_from_image(file_stream):
#         """Extract text from images using OCR"""
#         try:
#             image = Image.open(file_stream)
#             # Convert to RGB if needed
#             if image.mode != 'RGB':
#                 image = image.convert('RGB')
            
#             # Use pytesseract for OCR
#             text = pytesseract.image_to_string(image)
#             metadata = {
#                 'dimensions': image.size,
#                 'format': image.format,
#                 'mode': image.mode
#             }
#             return text, metadata
#         except Exception as e:
#             raise Exception(f"Image processing error: {str(e)}")
    
#     @staticmethod
#     def extract_text_from_docx(file_stream):
#         """Extract text from Word documents"""
#         try:
#             doc = Document(file_stream)
#             text = ""
#             for paragraph in doc.paragraphs:
#                 text += paragraph.text + "\n"
            
#             metadata = {
#                 'paragraphs': len(doc.paragraphs),
#                 'tables': len(doc.tables)
#             }
#             return text, metadata
#         except Exception as e:
#             raise Exception(f"DOCX processing error: {str(e)}")
    
#     @staticmethod
#     def extract_text_from_excel(file_stream):
#         """Extract text from Excel files"""
#         try:
#             workbook = openpyxl.load_workbook(file_stream)
#             text = ""
#             metadata = {'sheets': []}
            
#             for sheet_name in workbook.sheetnames:
#                 sheet = workbook[sheet_name]
#                 text += f"\n--- Sheet: {sheet_name} ---\n"
#                 metadata['sheets'].append(sheet_name)
                
#                 for row in sheet.iter_rows(values_only=True):
#                     row_text = " | ".join([str(cell) if cell is not None else "" for cell in row])
#                     if row_text.strip():
#                         text += row_text + "\n"
            
#             return text, metadata
#         except Exception as e:
#             raise Exception(f"Excel processing error: {str(e)}")
    
#     @staticmethod
#     def extract_text_from_csv(file_stream):
#         """Extract text from CSV files"""
#         try:
#             content = file_stream.read().decode('utf-8')
#             file_stream.seek(0)  # Reset stream position
            
#             csv_reader = csv.reader(io.StringIO(content))
#             text = ""
#             row_count = 0
            
#             for row in csv_reader:
#                 text += " | ".join(row) + "\n"
#                 row_count += 1
            
#             metadata = {'rows': row_count}
#             return text, metadata
#         except Exception as e:
#             raise Exception(f"CSV processing error: {str(e)}")

# def process_file(file):
#     """Process uploaded file based on its type"""
#     file_extension = Path(file.filename).suffix.lower()
#     file_stream = io.BytesIO(file.read())
    
#     processor = FileProcessor()
    
#     try:
#         if file_extension == '.pdf':
#             return processor.extract_text_from_pdf(file_stream)
#         elif file_extension in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff']:
#             return processor.extract_text_from_image(file_stream)
#         elif file_extension == '.docx':
#             return processor.extract_text_from_docx(file_stream)
#         elif file_extension in ['.xlsx', '.xls']:
#             return processor.extract_text_from_excel(file_stream)
#         elif file_extension == '.csv':
#             return processor.extract_text_from_csv(file_stream)
#         elif file_extension == '.txt':
#             content = file_stream.read().decode('utf-8')
#             return content, {'type': 'text', 'length': len(content)}
#         else:
#             raise Exception(f"Unsupported file type: {file_extension}")
#     except Exception as e:
#         raise Exception(f"Error processing {file.filename}: {str(e)}")

# def enhanced_groq_chat(messages, model="meta-llama/llama-4-scout-17b-16e-instruct", max_tokens=2048, temperature=0.7):
#     """Enhanced chat function with better parameters"""
#     completion = client.chat.completions.create(
#         model=model,
#         messages=messages,
#         temperature=temperature,
#         max_completion_tokens=max_tokens,
#         top_p=0.9,
#         stream=False,
#         stop=None,
#     )
#     return completion.choices[0].message.content

# def create_context_summary(file_contexts_dict):
#     """Create a summary of all uploaded files for context"""
#     if not file_contexts_dict:
#         return ""
    
#     context_summary = "\n=== UPLOADED FILES CONTEXT ===\n"
#     for file_id, context in file_contexts_dict.items():
#         context_summary += f"\nFile: {context['filename']}\n"
#         context_summary += f"Type: {context['file_type']}\n"
#         context_summary += f"Summary: {context.get('summary', 'No summary available')}\n"
#         context_summary += f"Key Content Preview: {context['content'][:500]}...\n"
#         context_summary += "---\n"
    
#     return context_summary

# @app.route('/upload-files', methods=['POST'])
# def upload_files():
#     """Enhanced endpoint to handle multiple file uploads"""
#     if 'files' not in request.files:
#         return jsonify({'error': 'No files uploaded'}), 400
    
#     files = request.files.getlist('files')
#     session_id = request.form.get('session_id', 'default')
    
#     if session_id not in file_contexts:
#         file_contexts[session_id] = {}
    
#     processed_files = []
    
#     for file in files:
#         if file.filename == '':
#             continue
            
#         try:
#             # Process the file
#             content, metadata = process_file(file)
            
#             # Generate file ID
#             file_id = hashlib.md5(f"{file.filename}{datetime.now().isoformat()}".encode()).hexdigest()[:8]
            
#             # Create enhanced summary based on file type
#             file_extension = Path(file.filename).suffix.lower()
            
#             if file_extension == '.pdf':
#                 summary_prompt = f"""Analyze this PDF document and provide a comprehensive summary including:
# 1. Main topic and purpose
# 2. Key concepts and themes
# 3. Important details and findings
# 4. Structure and organization
# 5. Any notable data or statistics

# Content (first 3000 chars): {content[:3000]}"""
            
#             elif file_extension in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff']:
#                 summary_prompt = f"""Analyze this image content extracted via OCR and provide:
# 1. What type of content this appears to be (document, chart, diagram, etc.)
# 2. Key information or text found
# 3. Context and purpose
# 4. Any important details or data

# OCR Content: {content[:2000]}"""
            
#             else:
#                 summary_prompt = f"""Analyze this document and provide a comprehensive summary including:
# 1. Document type and purpose
# 2. Main topics covered
# 3. Key information and insights
# 4. Structure and important sections

# Content (first 2500 chars): {content[:2500]}"""
            
#             # Generate summary
#             messages = [{"role": "user", "content": summary_prompt}]
#             summary = enhanced_groq_chat(messages, max_tokens=1024, temperature=0.5)
            
#             # Store file context
#             file_contexts[session_id][file_id] = {
#                 'filename': file.filename,
#                 'file_type': file_extension,
#                 'content': content,
#                 'metadata': metadata,
#                 'summary': summary,
#                 'upload_time': datetime.now().isoformat()
#             }
            
#             processed_files.append({
#                 'file_id': file_id,
#                 'filename': file.filename,
#                 'file_type': file_extension,
#                 'summary': summary,
#                 'metadata': metadata
#             })
            
#         except Exception as e:
#             processed_files.append({
#                 'filename': file.filename,
#                 'error': str(e)
#             })
    
#     return jsonify({
#         'processed_files': processed_files,
#         'session_id': session_id,
#         'total_files': len(file_contexts.get(session_id, {}))
#     })

# @app.route('/summarize-content', methods=['POST'])
# def summarize_content():
#     """Enhanced summarization with context awareness"""
#     data = request.json
#     session_id = data.get('session_id', 'default')
#     summary_type = data.get('type', 'comprehensive')  # comprehensive, brief, academic, technical
    
#     if session_id not in file_contexts or not file_contexts[session_id]:
#         return jsonify({'error': 'No files found for this session'}), 400
    
#     # Create context from all files
#     context_summary = create_context_summary(file_contexts[session_id])
    
#     # Generate different types of summaries
#     if summary_type == 'comprehensive':
#         prompt = f"""Based on all the uploaded files, provide a comprehensive summary that includes:

# 1. **Overview**: What are these documents about overall?
# 2. **Key Themes**: What are the main topics and themes across all files?
# 3. **Important Findings**: What are the most significant insights or data points?
# 4. **Connections**: How do these documents relate to each other?
# 5. **Actionable Insights**: What can be learned or applied from this content?

# {context_summary}

# Please provide a well-structured, detailed summary."""
    
#     elif summary_type == 'brief':
#         prompt = f"""Provide a brief, concise summary of the uploaded content focusing on:
# - Main topic(s)
# - Key points (3-5 bullet points)
# - Most important takeaway

# {context_summary}"""
    
#     elif summary_type == 'academic':
#         prompt = f"""Provide an academic-style analysis including:
# 1. Abstract/Executive Summary
# 2. Methodology (if applicable)
# 3. Key Findings
# 4. Conclusions
# 5. Implications for further study

# {context_summary}"""
    
#     else:  # technical
#         prompt = f"""Provide a technical summary focusing on:
# - Technical specifications or details
# - Data and metrics
# - Processes and procedures
# - Technical recommendations

# {context_summary}"""
    
#     try:
#         messages = [{"role": "user", "content": prompt}]
#         summary = enhanced_groq_chat(messages, max_tokens=2048, temperature=0.6)
        
#         return jsonify({
#             'summary': summary,
#             'summary_type': summary_type,
#             'files_count': len(file_contexts[session_id]),
#             'session_id': session_id
#         })
#     except Exception as e:
#         return jsonify({'error': str(e)}), 500

# @app.route('/chat', methods=['POST'])
# def chat():
#     """Enhanced chat with context awareness"""
#     data = request.json
#     user_message = data.get('message', '')
#     session_id = data.get('session_id', 'default')
    
#     if not user_message:
#         return jsonify({'error': 'No message provided'}), 400

#     # Initialize conversation context if not exists
#     if session_id not in conversation_contexts:
#         conversation_contexts[session_id] = []
    
#     # Add context from uploaded files if available
#     context_info = ""
#     if session_id in file_contexts and file_contexts[session_id]:
#         context_info = f"\n\n=== CONTEXT FROM UPLOADED FILES ===\n{create_context_summary(file_contexts[session_id])}\n=== END CONTEXT ===\n\n"
    
#     # Build conversation history
#     messages = []
    
#     # Add system message with context
#     system_message = f"""You are an advanced AI study assistant. You have access to uploaded documents and their content. 
# When answering questions, consider the context from uploaded files when relevant. 
# Provide detailed, helpful explanations and cite specific information from the uploaded content when applicable.
# Be educational, thorough, and adapt your explanations to the user's apparent level of understanding.

# {context_info}"""
    
#     if context_info:
#         messages.append({"role": "system", "content": system_message})
    
#     # Add recent conversation history (last 10 exchanges)
#     recent_context = conversation_contexts[session_id][-10:]
#     messages.extend(recent_context)
    
#     # Add current user message
#     messages.append({"role": "user", "content": user_message})
    
#     def generate():
#         try:
#             completion = client.chat.completions.create(
#                 model="meta-llama/llama-4-scout-17b-16e-instruct",
#                 messages=messages,
#                 temperature=0.7,
#                 max_completion_tokens=2048,
#                 top_p=0.9,
#                 stream=True,
#                 stop=None,
#             )
            
#             full_response = ""
#             for chunk in completion:
#                 content = chunk.choices[0].delta.content or ""
#                 full_response += content
#                 yield content
            
#             # Store conversation context
#             conversation_contexts[session_id].append({"role": "user", "content": user_message})
#             conversation_contexts[session_id].append({"role": "assistant", "content": full_response})
            
#         except Exception as e:
#             error_msg = f"\n\n**Error:** {str(e)}"
#             yield error_msg

#     return Response(stream_with_context(generate()), mimetype='text/plain')

# @app.route('/get-context', methods=['GET'])
# def get_context():
#     """Get current session context"""
#     session_id = request.args.get('session_id', 'default')
    
#     return jsonify({
#         'session_id': session_id,
#         'files': list(file_contexts.get(session_id, {}).keys()),
#         'file_details': [
#             {
#                 'file_id': fid,
#                 'filename': context['filename'],
#                 'file_type': context['file_type'],
#                 'upload_time': context['upload_time']
#             }
#             for fid, context in file_contexts.get(session_id, {}).items()
#         ],
#         'conversation_length': len(conversation_contexts.get(session_id, []))
#     })

# @app.route('/clear-context', methods=['POST'])
# def clear_context():
#     """Clear session context"""
#     data = request.json
#     session_id = data.get('session_id', 'default')
    
#     if session_id in file_contexts:
#         del file_contexts[session_id]
#     if session_id in conversation_contexts:
#         del conversation_contexts[session_id]
    
#     return jsonify({'message': 'Context cleared', 'session_id': session_id})

# # Legacy endpoint for backward compatibility
# @app.route('/summarize-pdf', methods=['POST'])
# def summarize_pdf():
#     """Legacy PDF summarization endpoint"""
#     if 'file' not in request.files:
#         return jsonify({'error': 'No file uploaded'}), 400
    
#     file = request.files['file']
    
#     try:
#         text, metadata = FileProcessor.extract_text_from_pdf(io.BytesIO(file.read()))
        
#         if not text.strip():
#             return jsonify({'error': 'No text found in PDF'}), 400

#         prompt = f"""Provide a comprehensive summary of this PDF document:

# **Document Info:**
# - Pages: {metadata.get('total_pages', 'Unknown')}
# - Title: {metadata.get('title', 'Unknown')}

# **Content Analysis:**
# Please summarize the key points, main themes, and important details from this document.

# **Content:**
# {text[:4000]}"""

#         messages = [{"role": "user", "content": prompt}]
#         summary = enhanced_groq_chat(messages, max_tokens=1536)
        
#         return jsonify({
#             'summary': summary,
#             'metadata': metadata
#         })
#     except Exception as e:
#         return jsonify({'error': str(e)}), 500

# if __name__ == '__main__':
#     # Create required directories
#     os.makedirs('uploads', exist_ok=True)
#     app.run(port=5001, debug=True)



# new updated
import os
import base64
import json
from datetime import datetime
from dotenv import load_dotenv
from flask import Flask, request, jsonify, Response, stream_with_context
from flask_cors import CORS
from groq import Groq
import io
import pytesseract
from PyPDF2 import PdfReader
from PIL import Image
from docx import Document
import openpyxl
import csv
from pathlib import Path
import hashlib
import logging

# Load environment variables
load_dotenv()

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Check for required environment variables
GROQ_API_KEY = os.environ.get("GROQ_API_KEY")
if not GROQ_API_KEY:
    raise ValueError("GROQ_API_KEY environment variable is not set. Please set it in your .env file.")

# Get port from environment (for Render deployment)
PORT = int(os.environ.get('PORT', 5001))

app = Flask(__name__)

# Configure CORS for production
CORS(app, resources={
    r"/*": {
        "origins": [
            "http://localhost:3000",  # React default port
            "http://localhost:5173",  # Vite default port
            "http://127.0.0.1:3000",
            "http://127.0.0.1:5173",
            "https://edu-hub-v1-f2ek.vercel.app",
            "https://*.onrender.com",  # Allow all Render subdomains
            "https://*.vercel.app",    # Allow Vercel domains
            # Add your actual frontend domain here
        ],
        "methods": ["GET", "POST", "PUT", "DELETE", "OPTIONS"],
        "allow_headers": ["Content-Type", "Authorization"]
    }
})

# Set file upload limits for free tier
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# Initialize Groq client with error handling
try:
    client = Groq(api_key=GROQ_API_KEY)
    logger.info("Groq client initialized successfully")
except Exception as e:
    logger.error(f"Error initializing Groq client: {str(e)}")
    raise

# Store conversation context and uploaded files
conversation_contexts = {}
file_contexts = {}

class FileProcessor:
    """Enhanced file processor supporting multiple formats"""
    
    @staticmethod
    def extract_text_from_pdf(file_stream):
        """Enhanced PDF text extraction with better formatting"""
        try:
            reader = PdfReader(file_stream)
            text = ""
            metadata = {
                'total_pages': len(reader.pages),
                'title': reader.metadata.get('/Title', 'Unknown') if reader.metadata else 'Unknown',
                'author': reader.metadata.get('/Author', 'Unknown') if reader.metadata else 'Unknown'
            }
            
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text() or ""
                if page_text.strip():
                    text += f"\n--- Page {i+1} ---\n{page_text}\n"
            
            return text, metadata
        except Exception as e:
            logger.error(f"PDF processing error: {str(e)}")
            raise Exception(f"PDF processing error: {str(e)}")
    
    @staticmethod
    def extract_text_from_image(file_stream):
        """Extract text from images using OCR"""
        try:
            image = Image.open(file_stream)
            # Convert to RGB if needed
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Use pytesseract for OCR
            text = pytesseract.image_to_string(image)
            metadata = {
                'dimensions': image.size,
                'format': image.format,
                'mode': image.mode
            }
            return text, metadata
        except Exception as e:
            logger.error(f"Image processing error: {str(e)}")
            raise Exception(f"Image processing error: {str(e)}")
    
    @staticmethod
    def extract_text_from_docx(file_stream):
        """Extract text from Word documents"""
        try:
            doc = Document(file_stream)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            metadata = {
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables)
            }
            return text, metadata
        except Exception as e:
            logger.error(f"DOCX processing error: {str(e)}")
            raise Exception(f"DOCX processing error: {str(e)}")
    
    @staticmethod
    def extract_text_from_excel(file_stream):
        """Extract text from Excel files"""
        try:
            workbook = openpyxl.load_workbook(file_stream)
            text = ""
            metadata = {'sheets': []}
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text += f"\n--- Sheet: {sheet_name} ---\n"
                metadata['sheets'].append(sheet_name)
                
                for row in sheet.iter_rows(values_only=True):
                    row_text = " | ".join([str(cell) if cell is not None else "" for cell in row])
                    if row_text.strip():
                        text += row_text + "\n"
            
            return text, metadata
        except Exception as e:
            logger.error(f"Excel processing error: {str(e)}")
            raise Exception(f"Excel processing error: {str(e)}")
    
    @staticmethod
    def extract_text_from_csv(file_stream):
        """Extract text from CSV files"""
        try:
            content = file_stream.read().decode('utf-8')
            file_stream.seek(0)  # Reset stream position
            
            csv_reader = csv.reader(io.StringIO(content))
            text = ""
            row_count = 0
            
            for row in csv_reader:
                text += " | ".join(row) + "\n"
                row_count += 1
            
            metadata = {'rows': row_count}
            return text, metadata
        except Exception as e:
            logger.error(f"CSV processing error: {str(e)}")
            raise Exception(f"CSV processing error: {str(e)}")

def process_file(file):
    """Process uploaded file based on its type"""
    file_extension = Path(file.filename).suffix.lower()
    file_stream = io.BytesIO(file.read())
    
    processor = FileProcessor()
    
    try:
        if file_extension == '.pdf':
            return processor.extract_text_from_pdf(file_stream)
        elif file_extension in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff']:
            return processor.extract_text_from_image(file_stream)
        elif file_extension == '.docx':
            return processor.extract_text_from_docx(file_stream)
        elif file_extension in ['.xlsx', '.xls']:
            return processor.extract_text_from_excel(file_stream)
        elif file_extension == '.csv':
            return processor.extract_text_from_csv(file_stream)
        elif file_extension == '.txt':
            content = file_stream.read().decode('utf-8')
            return content, {'type': 'text', 'length': len(content)}
        else:
            raise Exception(f"Unsupported file type: {file_extension}")
    except Exception as e:
        logger.error(f"Error processing {file.filename}: {str(e)}")
        raise Exception(f"Error processing {file.filename}: {str(e)}")

def enhanced_groq_chat(messages, model="meta-llama/llama-4-scout-17b-16e-instruct", max_tokens=2048, temperature=0.7):
    """Enhanced chat function with better parameters"""
    try:
        completion = client.chat.completions.create(
            model=model,
            messages=messages,
            temperature=temperature,
            max_completion_tokens=max_tokens,
            top_p=0.9,
            stream=False,
            stop=None,
        )
        return completion.choices[0].message.content
    except Exception as e:
        logger.error(f"Error in Groq chat: {str(e)}")
        raise

def create_context_summary(file_contexts_dict):
    """Create a summary of all uploaded files for context"""
    if not file_contexts_dict:
        return ""
    
    context_summary = "\n=== UPLOADED FILES CONTEXT ===\n"
    for file_id, context in file_contexts_dict.items():
        context_summary += f"\nFile: {context['filename']}\n"
        context_summary += f"Type: {context['file_type']}\n"
        context_summary += f"Summary: {context.get('summary', 'No summary available')}\n"
        context_summary += f"Key Content Preview: {context['content'][:500]}...\n"
        context_summary += "---\n"
    
    return context_summary

# Health check endpoint
@app.route('/health', methods=['GET'])
def health_check():
    """Health check endpoint for monitoring"""
    return jsonify({
        'status': 'healthy',
        'timestamp': datetime.now().isoformat(),
        'service': 'ai-tools-flask'
    })

# Root endpoint
@app.route('/', methods=['GET'])
def root():
    """Root endpoint"""
    return jsonify({
        'message': 'AI Tools Flask Server',
        'version': '1.0.0',
        'status': 'running',
        'endpoints': [
            '/health',
            '/upload-files',
            '/summarize-content',
            '/chat',
            '/get-context',
            '/clear-context'
        ]
    })

@app.route('/upload-files', methods=['POST'])
def upload_files():
    """Enhanced endpoint to handle multiple file uploads"""
    try:
        if 'files' not in request.files:
            return jsonify({'error': 'No files uploaded'}), 400
        
        files = request.files.getlist('files')
        session_id = request.form.get('session_id', 'default')
        
        if session_id not in file_contexts:
            file_contexts[session_id] = {}
        
        processed_files = []
        
        for file in files:
            if file.filename == '':
                continue
                
            try:
                # Process the file
                content, metadata = process_file(file)
                
                # Generate file ID
                file_id = hashlib.md5(f"{file.filename}{datetime.now().isoformat()}".encode()).hexdigest()[:8]
                
                # Create enhanced summary based on file type
                file_extension = Path(file.filename).suffix.lower()
                
                if file_extension == '.pdf':
                    summary_prompt = f"""Analyze this PDF document and provide a comprehensive summary including:
1. Main topic and purpose
2. Key concepts and themes
3. Important details and findings
4. Structure and organization
5. Any notable data or statistics

Content (first 3000 chars): {content[:3000]}"""
                
                elif file_extension in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff']:
                    summary_prompt = f"""Analyze this image content extracted via OCR and provide:
1. What type of content this appears to be (document, chart, diagram, etc.)
2. Key information or text found
3. Context and purpose
4. Any important details or data

OCR Content: {content[:2000]}"""
                
                else:
                    summary_prompt = f"""Analyze this document and provide a comprehensive summary including:
1. Document type and purpose
2. Main topics covered
3. Key information and insights
4. Structure and important sections

Content (first 2500 chars): {content[:2500]}"""
                
                # Generate summary
                messages = [{"role": "user", "content": summary_prompt}]
                summary = enhanced_groq_chat(messages, max_tokens=1024, temperature=0.5)
                
                # Store file context
                file_contexts[session_id][file_id] = {
                    'filename': file.filename,
                    'file_type': file_extension,
                    'content': content,
                    'metadata': metadata,
                    'summary': summary,
                    'upload_time': datetime.now().isoformat()
                }
                
                processed_files.append({
                    'file_id': file_id,
                    'filename': file.filename,
                    'file_type': file_extension,
                    'summary': summary,
                    'metadata': metadata
                })
                
                logger.info(f"Successfully processed file: {file.filename}")
                
            except Exception as e:
                logger.error(f"Error processing file {file.filename}: {str(e)}")
                processed_files.append({
                    'filename': file.filename,
                    'error': str(e)
                })
        
        return jsonify({
            'processed_files': processed_files,
            'session_id': session_id,
            'total_files': len(file_contexts.get(session_id, {}))
        })
    
    except Exception as e:
        logger.error(f"Error in upload_files: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/summarize-content', methods=['POST'])
def summarize_content():
    """Enhanced summarization with context awareness"""
    try:
        data = request.json
        session_id = data.get('session_id', 'default')
        summary_type = data.get('type', 'comprehensive')  # comprehensive, brief, academic, technical
        
        if session_id not in file_contexts or not file_contexts[session_id]:
            return jsonify({'error': 'No files found for this session'}), 400
        
        # Create context from all files
        context_summary = create_context_summary(file_contexts[session_id])
        
        # Generate different types of summaries
        if summary_type == 'comprehensive':
            prompt = f"""Based on all the uploaded files, provide a comprehensive summary that includes:

1. **Overview**: What are these documents about overall?
2. **Key Themes**: What are the main topics and themes across all files?
3. **Important Findings**: What are the most significant insights or data points?
4. **Connections**: How do these documents relate to each other?
5. **Actionable Insights**: What can be learned or applied from this content?

{context_summary}

Please provide a well-structured, detailed summary."""
        
        elif summary_type == 'brief':
            prompt = f"""Provide a brief, concise summary of the uploaded content focusing on:
- Main topic(s)
- Key points (3-5 bullet points)
- Most important takeaway

{context_summary}"""
        
        elif summary_type == 'academic':
            prompt = f"""Provide an academic-style analysis including:
1. Abstract/Executive Summary
2. Methodology (if applicable)
3. Key Findings
4. Conclusions
5. Implications for further study

{context_summary}"""
        
        else:  # technical
            prompt = f"""Provide a technical summary focusing on:
- Technical specifications or details
- Data and metrics
- Processes and procedures
- Technical recommendations

{context_summary}"""
        
        messages = [{"role": "user", "content": prompt}]
        summary = enhanced_groq_chat(messages, max_tokens=2048, temperature=0.6)
        
        return jsonify({
            'summary': summary,
            'summary_type': summary_type,
            'files_count': len(file_contexts[session_id]),
            'session_id': session_id
        })
    
    except Exception as e:
        logger.error(f"Error in summarize_content: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/chat', methods=['POST'])
def chat():
    """Enhanced chat with context awareness"""
    try:
        data = request.json
        user_message = data.get('message', '')
        session_id = data.get('session_id', 'default')
        
        if not user_message:
            return jsonify({'error': 'No message provided'}), 400

        # Initialize conversation context if not exists
        if session_id not in conversation_contexts:
            conversation_contexts[session_id] = []
        
        # Add context from uploaded files if available
        context_info = ""
        if session_id in file_contexts and file_contexts[session_id]:
            context_info = f"\n\n=== CONTEXT FROM UPLOADED FILES ===\n{create_context_summary(file_contexts[session_id])}\n=== END CONTEXT ===\n\n"
        
        # Build conversation history
        messages = []
        
        # Add system message with context
        system_message = f"""You are an advanced AI study assistant. You have access to uploaded documents and their content. 
When answering questions, consider the context from uploaded files when relevant. 
Provide detailed, helpful explanations and cite specific information from the uploaded content when applicable.
Be educational, thorough, and adapt your explanations to the user's apparent level of understanding.

{context_info}"""
        
        if context_info:
            messages.append({"role": "system", "content": system_message})
        
        # Add recent conversation history (last 10 exchanges)
        recent_context = conversation_contexts[session_id][-10:]
        messages.extend(recent_context)
        
        # Add current user message
        messages.append({"role": "user", "content": user_message})
        
        def generate():
            try:
                completion = client.chat.completions.create(
                    model="meta-llama/llama-4-scout-17b-16e-instruct",
                    messages=messages,
                    temperature=0.7,
                    max_completion_tokens=2048,
                    top_p=0.9,
                    stream=True,
                    stop=None,
                )
                
                full_response = ""
                for chunk in completion:
                    content = chunk.choices[0].delta.content or ""
                    full_response += content
                    yield content
                
                # Store conversation context
                conversation_contexts[session_id].append({"role": "user", "content": user_message})
                conversation_contexts[session_id].append({"role": "assistant", "content": full_response})
                
            except Exception as e:
                logger.error(f"Error in chat streaming: {str(e)}")
                error_msg = f"\n\n**Error:** {str(e)}"
                yield error_msg

        return Response(stream_with_context(generate()), mimetype='text/plain')
    
    except Exception as e:
        logger.error(f"Error in chat: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/get-context', methods=['GET'])
def get_context():
    """Get current session context"""
    try:
        session_id = request.args.get('session_id', 'default')
        
        return jsonify({
            'session_id': session_id,
            'files': list(file_contexts.get(session_id, {}).keys()),
            'file_details': [
                {
                    'file_id': fid,
                    'filename': context['filename'],
                    'file_type': context['file_type'],
                    'upload_time': context['upload_time']
                }
                for fid, context in file_contexts.get(session_id, {}).items()
            ],
            'conversation_length': len(conversation_contexts.get(session_id, []))
        })
    
    except Exception as e:
        logger.error(f"Error in get_context: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/clear-context', methods=['POST'])
def clear_context():
    """Clear session context"""
    try:
        data = request.json
        session_id = data.get('session_id', 'default')
        
        if session_id in file_contexts:
            del file_contexts[session_id]
        if session_id in conversation_contexts:
            del conversation_contexts[session_id]
        
        logger.info(f"Cleared context for session: {session_id}")
        return jsonify({'message': 'Context cleared', 'session_id': session_id})
    
    except Exception as e:
        logger.error(f"Error in clear_context: {str(e)}")
        return jsonify({'error': str(e)}), 500

# Legacy endpoint for backward compatibility
@app.route('/summarize-pdf', methods=['POST'])
def summarize_pdf():
    """Legacy PDF summarization endpoint"""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400
        
        file = request.files['file']
        
        text, metadata = FileProcessor.extract_text_from_pdf(io.BytesIO(file.read()))
        
        if not text.strip():
            return jsonify({'error': 'No text found in PDF'}), 400

        prompt = f"""Provide a comprehensive summary of this PDF document:

**Document Info:**
- Pages: {metadata.get('total_pages', 'Unknown')}
- Title: {metadata.get('title', 'Unknown')}

**Content Analysis:**
Please summarize the key points, main themes, and important details from this document.

**Content:**
{text[:4000]}"""

        messages = [{"role": "user", "content": prompt}]
        summary = enhanced_groq_chat(messages, max_tokens=1536)
        
        return jsonify({
            'summary': summary,
            'metadata': metadata
        })
    
    except Exception as e:
        logger.error(f"Error in summarize_pdf: {str(e)}")
        return jsonify({'error': str(e)}), 500

# Error handlers
@app.errorhandler(413)
def too_large(e):
    return jsonify({'error': 'File too large. Maximum size is 16MB.'}), 413

@app.errorhandler(404)
def not_found(e):
    return jsonify({'error': 'Endpoint not found'}), 404

@app.errorhandler(500)
def internal_error(e):
    logger.error(f"Internal server error: {str(e)}")
    return jsonify({'error': 'Internal server error'}), 500

if __name__ == '__main__':
    logger.info(f"Starting Flask app on port {PORT}")
    app.run(host='0.0.0.0', port=PORT, debug=False)